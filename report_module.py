# -*- coding: utf-8 -*-
"""
report_module.py
Формирование итогового отчёта аттестации в виде документа Word (.docx).

Требуется библиотека python-docx (pip install python-docx).
Если библиотека не установлена, модуль сохраняет отчёт в текстовый файл .txt,
чтобы результат тестирования не был потерян.
"""

import os
from datetime import datetime


def _safe_filename(name):
    """Убирает из ФИО символы, недопустимые в имени файла."""
    keep = "".join(ch if ch.isalnum() or ch in " _-" else "_" for ch in name)
    return "_".join(keep.split()) or "sotrudnik"


def save_report(name, dept, cattell_results, belbin_results, analyzer,
                cattell_test, belbin_test, folder="."):
    """
    Сохраняет отчёт. Возвращает (путь_к_файлу, предупреждение_или_None).

    name, dept       - данные сотрудника;
    cattell_results  - результат CattellTest.get_full_results();
    belbin_results   - словарь {роль: баллы};
    analyzer         - экземпляр ProfileAnalyzer;
    cattell_test     - экземпляр CattellTest;
    belbin_test      - экземпляр BelbinTest;
    folder           - папка для сохранения.
    """
    date_str = datetime.now().strftime("%Y-%m-%d_%H-%M")
    base = os.path.join(folder, f"Отчет_{_safe_filename(name)}_{date_str}")

    try:
        return _save_docx(base + ".docx", name, dept, cattell_results,
                          belbin_results, analyzer, cattell_test, belbin_test), None
    except ImportError:
        # python-docx не установлен — сохраняем текстовую версию
        path = base + ".txt"
        with open(path, "w", encoding="utf-8") as f:
            f.write(analyzer.analyze_intersections())
        warn = ("Библиотека python-docx не установлена, отчёт сохранён в текстовом виде.\n"
                "Для сохранения в Word выполните в командной строке:\n"
                "pip install python-docx")
        return path, warn


def _save_docx(path, name, dept, cattell_results, belbin_results,
               analyzer, cattell_test, belbin_test):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    stens = cattell_results['stens']
    doc = Document()

    # Базовый стиль
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Титульная часть ---
    title = doc.add_heading('Отчёт по результатам аттестации', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Комплексная оценка: опросник Р. Кеттелла 16PF (форма C) "
              "и опросник командных ролей Р. М. Белбина").italic = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run("Сотрудник: ").bold = True
    info.add_run(name)
    info2 = doc.add_paragraph()
    info2.add_run("Должность/отдел: ").bold = True
    info2.add_run(dept)
    info3 = doc.add_paragraph()
    info3.add_run("Дата тестирования: ").bold = True
    info3.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))

    # --- Достоверность ---
    doc.add_heading('Достоверность результатов', level=1)
    valid_p = doc.add_paragraph(cattell_results['validity_note'])
    if not cattell_results['valid']:
        valid_p.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        valid_p.runs[0].bold = True

    # --- Раздел I. Профиль Кеттелла ---
    doc.add_heading('I. Личностный профиль по Кеттеллу', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Фактор'
    hdr[1].text = 'Стен (1–10)'
    hdr[2].text = 'Интерпретация'
    for factor, sten in stens.items():
        row = table.add_row().cells
        row[0].text = cattell_test.factor_names[factor]
        row[1].text = str(sten)
        row[2].text = cattell_test.interpret_factor(factor, sten)

    # --- Раздел II. Роли Белбина ---
    doc.add_heading('II. Командные роли по Белбину', level=1)
    dominant = analyzer.get_dominant_roles()
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Роль'
    hdr2[1].text = 'Баллы (из 70)'
    hdr2[2].text = 'Статус'
    for role, score in sorted(belbin_results.items(), key=lambda x: -x[1]):
        row = table2.add_row().cells
        row[0].text = role
        row[1].text = str(score)
        row[2].text = 'Ведущая роль' if role in dominant else ''

    doc.add_paragraph()
    for role in dominant:
        p = doc.add_paragraph()
        p.add_run(f"«{role}». ").bold = True
        p.add_run(belbin_test.role_descriptions[role])

    # --- Раздел III. Сопоставление методик ---
    doc.add_heading('III. Сопоставление: подтверждаются ли роли личностным профилем', level=1)
    for role in dominant:
        status, text = analyzer.analyze_role(role)
        p = doc.add_paragraph()
        run = p.add_run(f"Роль «{role}» — {status}. ")
        run.bold = True
        if status == 'ПРОТИВОРЕЧИЕ':
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif status == 'ПОДТВЕРЖДЕНО':
            run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)
        p.add_run(text)

    # --- Раздел IV. Выводы ---
    doc.add_heading('IV. Общие выводы и рекомендации', level=1)
    for note in analyzer.general_conclusions():
        doc.add_paragraph(note, style='List Number')

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "Отчёт сформирован автоматически по результатам двух методик и носит "
        "рекомендательный характер. Кадровые решения должны приниматься с учётом "
        "интервью и рабочих показателей сотрудника."
    ).italic = True

    doc.save(path)
    return path


# --- Проверка модуля при автономном запуске ---
if __name__ == "__main__":
    from cattell_module import CattellTest
    from belbin_module import BelbinTest
    from analyzer import ProfileAnalyzer

    cattell = CattellTest()
    belbin = BelbinTest()

    mock_cattell = cattell.get_full_results({q: 'a' for q in range(1, 106)})
    mock_belbin = belbin.calculate_results({b: [2, 2, 2, 1, 1, 1, 1, 0] for b in range(1, 8)})
    analyzer = ProfileAnalyzer("Тест Тестович", "Отдел контроля качества",
                               mock_belbin, mock_cattell,
                               belbin_test=belbin, cattell_test=cattell)

    path, warn = save_report("Тест Тестович", "Отдел контроля качества",
                             mock_cattell, mock_belbin, analyzer, cattell, belbin)
    print("Отчёт сохранён:", path)
    if warn:
        print(warn)
